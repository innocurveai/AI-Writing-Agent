"""
고유 문체 기반 AI Writing Agent (Streamlit)

Supabase(pgvector)에 아래 SQL을 적용한 뒤 사용하세요.
- user_documents: id, filename, file_type, upload_date, genre
- document_embeddings: id, doc_id, chunk_content, embedding_vector(vector(1536))
- RPC: match_document_embeddings (`supabase_schema.sql` 참고)
"""

from __future__ import annotations

import base64
import functools
import io
import os
import re
import uuid
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import pdfplumber
import streamlit as st
import streamlit.components.v1 as components
from dotenv import dotenv_values, load_dotenv
from langchain_text_splitters import RecursiveCharacterTextSplitter
from openai import OpenAI
from supabase import Client, create_client

# `streamlit run` 시 cwd가 프로젝트 루트가 아닐 수 있어, .env는 이 파일 기준으로 찾음
_APP_DIR = Path(__file__).resolve().parent

# 스타일 레퍼런스 벡터 검색 시 가져올 상위 청크 수(슬라이더 제거 후 고정)
STYLE_REFERENCE_TOP_K = 6


def _truncate_label(name: str, max_len: int = 50) -> str:
    n = (name or "").strip()
    if len(n) <= max_len:
        return n
    return n[: max_len - 1] + "…"


def _sanitize_docx_plain_text(s: str) -> str:
    """OOXML w:t 에서 문제 되는 제어 문자만 제거."""
    return re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", s)


def _draft_docx_ooxml_zip_bytes(body: str) -> bytes:
    """
    python-docx 없이 최소 OOXML(.docx)만 생성.
    Streamlit Cloud에서 패키지 미설치·import 실패여도 Word에서 열 수 있게 함.
    """
    from xml.sax.saxutils import escape

    lines = body.splitlines()
    if not lines:
        lines = [body] if body else [""]

    paras: list[str] = []
    for line in lines:
        t = escape(_sanitize_docx_plain_text(line))
        paras.append(
            f'<w:p><w:r><w:t xml:space="preserve">{t}</w:t></w:r></w:p>'
        )
    paras_xml = "".join(paras)

    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"<w:body>{paras_xml}"
        "<w:sectPr>"
        '<w:pgSz w:w="11906" w:h="16838"/>'
        '<w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" '
        'w:header="708" w:footer="708" w:gutter="0"/>'
        "</w:sectPr></w:body></w:document>"
    )

    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""

    package_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", package_rels)
        zf.writestr("word/document.xml", document_xml.encode("utf-8"))
    buf.seek(0)
    return buf.getvalue()


def build_draft_docx_bytes(body: str) -> bytes:
    """생성 본문을 줄 단위 Word(.docx)로 직렬화. python-docx가 없으면 OOXML(zip) 폴백."""
    try:
        from docx import Document

        doc = Document()
        doc.core_properties.title = "생성 초안"
        if not body.strip():
            doc.add_paragraph("")
        else:
            for line in body.splitlines():
                doc.add_paragraph(line)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
    except Exception:
        return _draft_docx_ooxml_zip_bytes(body)


def _short_upload_ts(val: Any) -> str:
    if val is None:
        return ""
    s = str(val).strip().replace("T", " ", 1)
    return s[:16] if len(s) > 16 else s


# ---------------------------------------------------------------------------
# 한지 톤 UI + 상단 기와(실사 이미지 또는 CSS 폴백)
# ---------------------------------------------------------------------------


@functools.lru_cache(maxsize=4)
def _giwa_header_jpeg_data_url(path_str: str, mtime_ns: int) -> str | None:
    """static/giwa_header.png → 리사이즈 JPEG data URL (헤더용, 캐시)."""
    path = Path(path_str)
    if not path.is_file():
        return None
    try:
        from io import BytesIO

        from PIL import Image

        raw = path.read_bytes()
        im = Image.open(BytesIO(raw))
        if im.mode in ("RGBA", "P"):
            im = im.convert("RGB")
        w, h = im.size
        max_w = 1600
        if w > max_w:
            nh = max(1, int(h * (max_w / w)))
            im = im.resize((max_w, nh), Image.Resampling.LANCZOS)
        buf = BytesIO()
        im.save(buf, format="JPEG", quality=84, optimize=True)
        b64 = base64.b64encode(buf.getvalue()).decode("ascii")
        return f"data:image/jpeg;base64,{b64}"
    except Exception:
        try:
            b64 = base64.b64encode(path.read_bytes()).decode("ascii")
            return f"data:image/png;base64,{b64}"
        except Exception:
            return None


def _giwa_header_background_url() -> str | None:
    p = _APP_DIR / "static" / "giwa_header.png"
    if not p.is_file():
        return None
    try:
        mt = int(p.stat().st_mtime_ns)
    except AttributeError:
        mt = int(p.stat().st_mtime)
    return _giwa_header_jpeg_data_url(str(p.resolve()), mt)


def inject_hanji_theme() -> None:
    giwa_url = _giwa_header_background_url()
    if giwa_url:
        st_header_core = """
  [data-testid="stHeader"] {
    position: relative !important;
    z-index: 1002;
    min-height: 3.25rem;
    background-color: #252628 !important;
    background-image:
      linear-gradient(
        180deg,
        rgba(70, 72, 74, 0.15) 0%,
        rgba(25, 26, 28, 0.35) 40%,
        rgba(6, 6, 8, 0.82) 100%
      ),
      linear-gradient(
        90deg,
        rgba(0, 0, 0, 0.35) 0%,
        transparent 18%,
        transparent 82%,
        rgba(0, 0, 0, 0.35) 100%
      ),
      url("__GIWA_URL__");
    background-size: 100% 100%, 100% 100%, cover;
    background-position: center, center, center 42%;
    background-repeat: no-repeat;
    background-blend-mode: soft-light, multiply, normal;
    border-radius: 0 0 16px 16px;
    border-bottom: 4px solid #070708;
    box-shadow:
      inset 0 0 24px rgba(0, 0, 0, 0.45),
      inset 0 -16px 40px rgba(0, 0, 0, 0.35),
      0 8px 22px rgba(0, 0, 0, 0.38);
    backdrop-filter: none;
    color: #fdf6e8 !important;
  }
  [data-testid="stHeader"]::before {
    content: "";
    pointer-events: none;
    position: absolute;
    inset: 0;
    border-radius: 0 0 16px 16px;
    background: radial-gradient(
      120% 80% at 50% -10%,
      rgba(255, 255, 255, 0.12) 0%,
      transparent 55%
    );
    mix-blend-mode: overlay;
    opacity: 0.55;
  }
"""
        st_header_core = st_header_core.replace("__GIWA_URL__", giwa_url)
    else:
        st_header_core = """
  [data-testid="stHeader"] {
    position: relative !important;
    z-index: 1002;
    background-color: #232224 !important;
    background-image:
      linear-gradient(
        180deg,
        rgba(120, 118, 115, 0.45) 0%,
        rgba(40, 38, 37, 0.1) 38%,
        rgba(0, 0, 0, 0.55) 100%
      ),
      repeating-linear-gradient(
        92deg,
        #1c1b1a 0 3px,
        #2f2d2e 3px 22px,
        #252324 22px 24px,
        #3d3b3c 24px 26px,
        #262424 26px 45px,
        #181716 45px 47px,
        #343233 47px 49px,
        #222020 49px 68px,
        #141312 68px 70px,
        #2a2829 70px 88px
      ),
      repeating-linear-gradient(
        -6deg,
        transparent 0 18px,
        rgba(0, 0, 0, 0.35) 18px 20px,
        transparent 20px 52px
      ),
      repeating-linear-gradient(
        90deg,
        rgba(255, 255, 255, 0) 0 36px,
        rgba(255, 255, 255, 0.07) 36px 38px,
        rgba(255, 255, 255, 0) 38px 88px
      );
    background-size: 100% 100%, 88px 140%, 100% 40px, 88px 100%;
    background-position: 0 0, -6px -32px, 0 0, 0 0;
    background-repeat: no-repeat, repeat, repeat, repeat;
    border-radius: 0 0 14px 14px;
    border-bottom: 4px solid #0a0a0b;
    box-shadow:
      inset 0 1px 0 rgba(255, 255, 255, 0.1),
      inset 0 -10px 28px rgba(0, 0, 0, 0.65),
      0 8px 18px rgba(0, 0, 0, 0.28);
    backdrop-filter: none;
    color: #f6ecd8 !important;
  }
  [data-testid="stHeader"]::before {
    content: "";
    pointer-events: none;
    position: absolute;
    inset: 0 0 auto 0;
    height: 46%;
    border-radius: 0 0 14px 14px;
    background: repeating-linear-gradient(
      90deg,
      rgba(255, 255, 255, 0) 0 40px,
      rgba(255, 255, 255, 0.09) 40px 42px,
      rgba(255, 255, 255, 0) 42px 86px
    );
    mix-blend-mode: soft-light;
    opacity: 0.85;
  }
"""

    st.markdown(
        """
<style>
  /* OS/브라우저 다크 선호와 무관하게 앱 전역을 라이트 톤으로 고정 */
  :root {
    color-scheme: light !important;
  }
  html, body, [data-testid="stAppViewContainer"] {
    color-scheme: light !important;
    background-color: #e8ddc4 !important;
    background-image:
      repeating-linear-gradient(
        0deg,
        rgba(255, 255, 255, 0.04) 0 2px,
        transparent 2px 5px
      ),
      repeating-linear-gradient(
        90deg,
        rgba(160, 130, 80, 0.035) 0 1px,
        transparent 1px 88px
      ),
      radial-gradient(ellipse at 20% 0%, rgba(255, 248, 220, 0.45), transparent 55%),
      radial-gradient(ellipse at 80% 100%, rgba(210, 185, 130, 0.2), transparent 50%);
    background-repeat: repeat, repeat, no-repeat, no-repeat;
    background-attachment: scroll;
    background-blend-mode: normal, normal, soft-light, multiply;
  }
  /* 문서 전체가 세로로 늘어나게 — stMain에 overflow:visible 강제하면 내부 스크롤·휠이 먹통 나는 경우 있음 */
  html, body {
    overflow-x: hidden !important;
    overflow-y: auto !important;
    height: auto !important;
    max-height: none !important;
  }
  .stApp {
    background: transparent !important;
    height: auto !important;
    min-height: 100vh !important;
    max-height: none !important;
    overflow-x: hidden !important;
    overflow-y: visible !important;
  }
  [data-testid="stAppViewContainer"] {
    min-height: 100vh !important;
    height: auto !important;
    max-height: none !important;
    overflow: visible !important;
  }
  /* wide 루트 행: 높이만 자연스럽게 (display:flex는 덮어쓰지 않음) */
  [data-testid="stAppViewContainer"] > div {
    height: auto !important;
    max-height: none !important;
    overflow: visible !important;
  }
  /* 메인 전 세로: .block-container 밖·아래 구간도 Streamlit 기본 흰색이 안 보이게 한지 톤 통일 */
  section[data-testid="stMain"] {
    height: auto !important;
    max-height: none !important;
    background-color: #e8ddc4 !important;
    background-image:
      repeating-linear-gradient(
        0deg,
        rgba(255, 255, 255, 0.04) 0 2px,
        transparent 2px 5px
      ),
      repeating-linear-gradient(
        90deg,
        rgba(160, 130, 80, 0.035) 0 1px,
        transparent 1px 88px
      ),
      radial-gradient(ellipse at 20% 0%, rgba(255, 248, 220, 0.45), transparent 55%),
      radial-gradient(ellipse at 80% 100%, rgba(210, 185, 130, 0.2), transparent 50%) !important;
    background-repeat: repeat, repeat, no-repeat, no-repeat !important;
    background-attachment: scroll !important;
    background-blend-mode: normal, normal, soft-light, multiply !important;
  }
  section[data-testid="stMain"] > div {
    background-color: transparent !important;
    background-image: none !important;
  }
  .stAlert,
  div[data-testid="stNotification"] {
    overflow: visible !important;
    max-height: none !important;
  }
"""
        + st_header_core
        + """
  [data-testid="stHeader"] a,
  [data-testid="stHeader"] button,
  [data-testid="stHeader"] p,
  [data-testid="stHeader"] span,
  [data-testid="stHeader"] label {
    color: #f8f0dc !important;
    text-shadow:
      0 1px 2px rgba(0, 0, 0, 0.95),
      0 0 12px rgba(0, 0, 0, 0.55);
  }
  [data-testid="stHeader"] svg {
    fill: #f4e8cc !important;
    filter: drop-shadow(0 1px 1px rgba(0, 0, 0, 0.9));
  }
  /* 사이드바: 베이스 톤 + 그라데이션(아래로 스크롤해도 메인과 색 끊김 완화) */
  [data-testid="stSidebar"] {
    align-self: stretch !important;
    box-sizing: border-box !important;
    background-color: #e4dac0 !important;
    background-image: linear-gradient(
      175deg,
      rgba(244, 234, 210, 0.92) 0%,
      rgba(220, 202, 168, 0.88) 100%
    ) !important;
    background-repeat: no-repeat !important;
    background-size: 100% 100% !important;
    border-right: 1px solid rgba(140, 118, 78, 0.18);
  }
  [data-testid="stSidebar"] > div {
    background-color: transparent !important;
    background-image: none !important;
    min-height: unset !important;
    height: auto !important;
  }
  .block-container {
    background: rgba(252, 248, 236, 0.55) !important;
    border-radius: 10px;
    padding-top: 1.5rem;
    padding-bottom: 2rem;
    box-shadow: inset 0 0 0 1px rgba(120, 98, 60, 0.08);
  }
  div[data-testid="stExpander"] details {
    background: rgba(252, 246, 230, 0.75) !important;
    border: 1px solid rgba(130, 110, 70, 0.12);
  }
  .stAlert, [data-baseweb="notification"] {
    background: rgba(255, 251, 235, 0.95) !important;
  }
  .stApp,
  header[data-testid="stHeader"],
  section[data-testid="stSidebar"] {
    color-scheme: light !important;
  }
</style>
""",
        unsafe_allow_html=True,
    )


def _request_scroll_to_bottom() -> None:
    """main() 끝에서 스크롤하도록 플래그만 설정(생성 결과 DOM이 그려진 뒤 실행)."""
    st.session_state["_pending_scroll_bottom"] = True


def _scroll_to_show_new_content() -> None:
    """생성 결과·에러가 뷰포트 아래일 때 스크롤(Streamlit iframe·부모 창 모두 시도)."""
    components.html(
        """
<script>
(function () {
  /* 문서(scrollingElement)만 스크롤. stMain.scrollTop·scrollIntoView는
     내부 스크롤에 갇혀 위아래 휠이 안 먹는 현상을 유발할 수 있음 */
  function scrollDocRoot(w) {
    if (!w || !w.document) return;
    var se = w.document.scrollingElement || w.document.documentElement;
    if (se) {
      se.scrollTop = se.scrollHeight;
    }
  }
  function run() {
    scrollDocRoot(window.parent);
    if (window.parent && window.parent.parent) {
      scrollDocRoot(window.parent.parent);
    }
    scrollDocRoot(window);
  }
  run();
  setTimeout(run, 80);
  setTimeout(run, 300);
})();
</script>
        """,
        height=0,
        width=0,
    )


# ---------------------------------------------------------------------------
# 설정 로드 (Streamlit Cloud: st.secrets 우선, 로컬: .env / 환경변수)
# ---------------------------------------------------------------------------


def _load_dotenv_if_present() -> None:
    # 프로젝트 루트 .env는 **항상 우선** (override=True)
    # Streamlit/쉘이 SUPABASE_URL="" 처럼 빈 값을 넣어두면 override=False일 때 .env가 무시됨
    enc = "utf-8-sig"  # BOM·한글 주석 호환
    load_dotenv(_APP_DIR / ".env", override=True, encoding=enc)
    # cwd의 .env는 없는 키만 채움
    load_dotenv(Path.cwd() / ".env", override=False, encoding=enc)


def get_config() -> dict[str, str | None]:
    _load_dotenv_if_present()
    out: dict[str, str | None] = {
        "SUPABASE_URL": None,
        "SUPABASE_ANON_KEY": None,
        "OPENAI_API_KEY": None,
    }
    for key in out:
        secret_val: str | None = None
        try:
            if hasattr(st, "secrets") and key in st.secrets:
                raw = st.secrets[key]
                if raw is not None and str(raw).strip():
                    secret_val = str(raw).strip()
        except (FileNotFoundError, RuntimeError, KeyError):
            pass
        if secret_val:
            out[key] = secret_val
            continue
        env_val = os.getenv(key)
        out[key] = env_val.strip() if env_val and env_val.strip() else None
    return out


def missing_keys(cfg: dict[str, str | None]) -> list[str]:
    return [k for k, v in cfg.items() if not v]


# ---------------------------------------------------------------------------
# Supabase / OpenAI 클라이언트
# ---------------------------------------------------------------------------


@st.cache_resource
def supabase_client(url: str, key: str) -> Client:
    return create_client(url, key)


@st.cache_resource
def openai_client(api_key: str) -> OpenAI:
    return OpenAI(api_key=api_key)


# ---------------------------------------------------------------------------
# 텍스트 추출 · 청킹 · 임베딩
# ---------------------------------------------------------------------------


def extract_text_from_upload(name: str, data: bytes) -> str:
    """AI/앱이 읽을 수 있는 평문·PDF만 지원."""
    lower = name.lower()
    if lower.endswith((".txt", ".md", ".markdown")):
        return data.decode("utf-8", errors="replace")
    if lower.endswith(".json"):
        return data.decode("utf-8", errors="replace")
    if lower.endswith(".pdf"):
        text_parts: list[str] = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                if t.strip():
                    text_parts.append(t)
        return "\n\n".join(text_parts)
    raise ValueError("지원 형식: .pdf, .txt, .md, .json")


def chunk_text(text: str) -> list[str]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=900,
        chunk_overlap=140,
        length_function=len,
    )
    chunks = splitter.split_text(text)
    return [c.strip() for c in chunks if c.strip()]


def embed_texts(client: OpenAI, texts: list[str]) -> list[list[float]]:
    if not texts:
        return []
    resp = client.embeddings.create(
        model="text-embedding-3-small",
        input=texts,
    )
    # 응답 순서 보존
    by_index = sorted(resp.data, key=lambda d: d.index)
    return [d.embedding for d in by_index]


# ---------------------------------------------------------------------------
# DB 작업
# ---------------------------------------------------------------------------


def insert_document(
    sb: Client,
    filename: str,
    file_type: str,
    genre: str | None,
) -> str:
    row = {
        "id": str(uuid.uuid4()),
        "filename": filename,
        "file_type": file_type,
        "upload_date": datetime.now(timezone.utc).isoformat(),
        "genre": genre,
    }
    sb.table("user_documents").insert(row).execute()
    return row["id"]


def insert_embeddings(
    sb: Client,
    doc_id: str,
    chunks: list[str],
    vectors: list[list[float]],
) -> None:
    rows: list[dict[str, Any]] = []
    for content, vec in zip(chunks, vectors, strict=True):
        rows.append(
            {
                "id": str(uuid.uuid4()),
                "doc_id": doc_id,
                "chunk_content": content,
                "embedding_vector": vec,
            }
        )
    if rows:
        sb.table("document_embeddings").insert(rows).execute()


def list_user_documents(sb: Client) -> list[dict[str, Any]]:
    res = (
        sb.table("user_documents")
        .select("id, filename, file_type, upload_date, genre")
        .order("upload_date", desc=True)
        .execute()
    )
    return list(res.data or [])


def delete_document(sb: Client, doc_id: str) -> None:
    sb.table("user_documents").delete().eq("id", doc_id).execute()


def fetch_chunks_fallback(
    sb: Client,
    limit: int,
    filter_doc_ids: list[str] | None,
) -> list[dict[str, Any]]:
    """RPC가 0건일 때 저장된 청크를 직접 읽어 레퍼런스로 사용(유사도 없음)."""
    lim = max(1, min(limit * 2, 48))
    q = sb.table("document_embeddings").select("id, doc_id, chunk_content").limit(lim)
    if filter_doc_ids:
        q = q.in_("doc_id", filter_doc_ids)
    res = q.execute()
    rows = list(res.data or [])
    out: list[dict[str, Any]] = []
    for row in rows[:limit]:
        out.append(
            {
                "id": row.get("id"),
                "doc_id": row.get("doc_id"),
                "chunk_content": row.get("chunk_content") or "",
            }
        )
    return out


def match_chunks(
    sb: Client,
    client: OpenAI,
    query: str,
    match_count: int,
    match_threshold: float,
    filter_doc_ids: list[str] | None,
    filter_genre: str | None,
) -> list[dict[str, Any]]:
    """키워드 임베딩으로 RPC 검색. 유사도(임계값) 때문에 0건이면 임계값 0으로 한 번 더 시도."""
    qvec = embed_texts(client, [query])
    if not qvec:
        return []

    def _call_rpc(th: float) -> list[dict[str, Any]]:
        payload: dict[str, Any] = {
            "query_embedding": qvec[0],
            "match_count": max(1, match_count),
            "match_threshold": th,
        }
        if filter_doc_ids is not None:
            payload["filter_doc_ids"] = filter_doc_ids
        if filter_genre is not None:
            payload["filter_genre"] = filter_genre
        res = sb.rpc("match_document_embeddings", payload).execute()
        return list(res.data or [])

    try:
        rows = _call_rpc(match_threshold)
        if not rows and match_threshold > 0.001:
            rows = _call_rpc(0.0)
        if not rows:
            rows = fetch_chunks_fallback(sb, max(1, match_count), filter_doc_ids)
        return rows
    except Exception as e:  # noqa: BLE001
        raise RuntimeError(
            "벡터 검색(RPC match_document_embeddings)에 실패했습니다. "
            "Supabase에서 `supabase_schema.sql`이 적용됐는지, 컬럼명이 `embedding_vector`인지 확인하세요. "
            f"원인: {e}"
        ) from e


# ---------------------------------------------------------------------------
# LLM 창작 (문체 레퍼런스 주입)
# ---------------------------------------------------------------------------


GENRE_LABELS = {"poetry": "시", "essay": "에세이", "novel": "소설"}


def get_chat_model() -> str:
    """`.env`의 OPENAI_CHAT_MODEL(예: gpt-4o-mini). 없으면 gpt-5.4-mini."""
    env_path = _APP_DIR / ".env"
    if env_path.is_file():
        load_dotenv(env_path, override=True)
    return (os.getenv("OPENAI_CHAT_MODEL") or "gpt-5.4-mini").strip()


def _normalize_chat_message_content(raw: Any) -> str:
    """Chat Completions의 message.content가 str 또는 파트 리스트일 때 본문만 합침."""
    if raw is None:
        return ""
    if isinstance(raw, str):
        return raw
    if isinstance(raw, list):
        parts: list[str] = []
        for block in raw:
            if isinstance(block, dict):
                if block.get("type") == "text":
                    parts.append(str(block.get("text") or ""))
                continue
            btype = getattr(block, "type", None)
            if btype == "text" or btype is None:
                tx = getattr(block, "text", None)
                if tx is not None:
                    parts.append(str(tx))
        return "".join(parts)
    return str(raw)


def build_style_block(chunks: list[dict[str, Any]]) -> str:
    parts: list[str] = []
    for i, row in enumerate(chunks, start=1):
        content = row.get("chunk_content") or ""
        parts.append(f"[레퍼런스 {i}]\n{content.strip()}")
    return "\n\n".join(parts)


def llm_draft(
    oai: OpenAI,
    genre_key: str,
    keywords: str,
    style_block: str,
    plot_outline: str | None,
) -> str:
    genre_ko = GENRE_LABELS.get(genre_key, genre_key)
    plot_section = ""
    if plot_outline:
        plot_section = (
            "\n\n[기획(플롯) — 초안은 반드시 이 구성에 맞출 것]\n" + plot_outline.strip()
        )

    system = (
        "당신은 한국어 창작 보조 작가입니다. "
        "기계적이거나 설명조의 말투, 메타 코멘트(예: '다음은 초안입니다')는 금지합니다. "
        "반드시 아래 '스타일 레퍼런스'에서 보이는 어조, 단어 선택, 문장 리듬, "
        "문장 길이의 분포를 모방하여 같은 장르의 새 작품 초안만 출력합니다."
    )
    user = (
        f"장르: {genre_ko}\n"
        f"키워드/주제: {keywords}\n"
        f"{plot_section}\n\n"
        f"스타일 레퍼런스(저자 작품에서 발췌):\n{style_block}\n\n"
        f"위 조건을 만족하는 {genre_ko} 초안을 작성하세요. 제목이 필요하면 자연스럽게 포함해도 됩니다."
    )
    comp = oai.chat.completions.create(
        model=get_chat_model(),
        temperature=0.9,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
    )
    if not comp.choices:
        return ""
    choice0 = comp.choices[0]
    msg = choice0.message
    refusal = getattr(msg, "refusal", None)
    body = _normalize_chat_message_content(getattr(msg, "content", None)).strip()
    if refusal and not body:
        raise RuntimeError(
            f"모델이 본문 생성을 거절했습니다: {refusal}. "
            "프롬프트·키워드를 바꾸거나 OPENAI_CHAT_MODEL을 다른 모델로 지정해 보세요."
        )
    if not body and getattr(choice0, "finish_reason", None) == "content_filter":
        raise RuntimeError("콘텐츠 필터로 인해 응답이 비었습니다. 키워드·레퍼런스를 조정해 보세요.")
    return body


# ---------------------------------------------------------------------------
# Streamlit UI
# ---------------------------------------------------------------------------


def render_sidebar_docs(
    sb: Client | None,
    cfg_ok: bool,
    cfg: dict[str, str | None],
) -> None:
    st.sidebar.header("학습 데이터 관리")
    if not cfg_ok or sb is None:
        st.sidebar.info("키가 준비되면 파일을 업로드할 수 있습니다.")
        return

    st.sidebar.caption(
        "업로드한 모든 문서는 **장르 구분 없이** 학습되며, "
        "아래에서 고른 시·에세이·소설 생성 시 함께 검색됩니다."
    )

    up = st.sidebar.file_uploader(
        "추가 파일 업로드 (PDF / TXT / MD / JSON)",
        type=["pdf", "txt", "md", "json"],
        accept_multiple_files=False,
    )
    if up is not None:
        if st.sidebar.button("이 파일 학습하기", type="primary"):
            try:
                raw = up.getvalue()
                text = extract_text_from_upload(up.name, raw)
                if not text.strip():
                    st.sidebar.error("추출된 텍스트가 비어 있습니다. 다른 PDF를 시도해 보세요.")
                else:
                    chunks = chunk_text(text)
                    oai = openai_client(str(cfg.get("OPENAI_API_KEY") or ""))
                    vectors = embed_texts(oai, chunks)
                    doc_id = insert_document(sb, up.name, os.path.splitext(up.name)[1].lower(), None)
                    insert_embeddings(sb, doc_id, chunks, vectors)
                    st.sidebar.success(f"학습 완료: {up.name} ({len(chunks)}청크)")
                    st.rerun()
            except Exception as e:  # noqa: BLE001
                st.sidebar.error(f"업로드 처리 실패: {e}")

    try:
        docs = list_user_documents(sb)
    except Exception as e:  # noqa: BLE001
        st.sidebar.error(f"문서 목록 조회 실패: {e}")
        st.sidebar.code(
            "Supabase 테이블·RPC(match_document_embeddings) 설정을 확인하세요.",
            language="text",
        )
        return

    if not docs:
        st.sidebar.caption("아직 학습된 문서가 없습니다.")
        return

    st.sidebar.subheader(f"저장된 문서 ({len(docs)}건)")
    st.sidebar.caption("파일 목록은 **문서 삭제**를 펼쳐 선택할 수 있습니다.")

    # expander 안에서는 st.sidebar.* 가 아니라 st.* 를 써야 블록 안에 붙음
    with st.sidebar.expander("문서 삭제", expanded=False):
        pick = st.selectbox(
            "삭제할 문서",
            options=list(range(len(docs))),
            format_func=lambda i: (
                f"{i + 1}. {_truncate_label(str(docs[i].get('filename') or ''), 40)}"
                f" · {_short_upload_ts(docs[i].get('upload_date'))}"
            ),
            key="sidebar_doc_delete_index",
            label_visibility="collapsed",
        )
        if st.button("선택한 문서 삭제", type="secondary", key="sidebar_doc_delete_btn"):
            try:
                delete_document(sb, str(docs[pick]["id"]))
                st.success("삭제되었습니다.")
                st.rerun()
            except Exception as ex:  # noqa: BLE001
                st.error(str(ex))


def main() -> None:
    st.set_page_config(
        page_title="고유 문체 기반 AI Writing Agent",
        layout="wide",
    )
    inject_hanji_theme()
    st.title("고유 문체 기반 AI Writing Agent")
    st.caption("시 · 에세이 · 소설 — 업로드한 작품 문체를 학습해 결과를 만듭니다.")

    cfg = get_config()
    miss = missing_keys(cfg)
    cfg_ok = not miss

    if miss:
        env_path = _APP_DIR / ".env"
        st.warning(
            "다음 키가 필요합니다: **"
            + ", ".join(miss)
            + "**\n\n"
            "- **로컬**: `"
            + str(env_path)
            + "` 파일에 위 변수를 넣고 저장한 뒤 앱을 새로고침하세요. "
            "(저장 위치는 이 앱 파일과 **같은 폴더**여야 합니다.)\n"
            "- **Streamlit Cloud**: Git에 올라가지 않는 `.env`는 배포 서버에 **없습니다**. "
            "앱 관리 화면 → **Settings → Secrets**에 동일 이름으로 넣어야 합니다."
        )
        # 디스크에 있는 .env만 읽음(저장 여부 진단). 에디터 미저장 버퍼는 여기 반영 안 됨
        if env_path.is_file():
            dv = dotenv_values(env_path, encoding="utf-8-sig")
            disk_blank = [k for k in miss if not str(dv.get(k) or "").strip()]
            if disk_blank == miss:
                st.error(
                    "**`.env`는 있지만 위 키 값이 디스크 기준으로 비어 있습니다.** "
                    "에디터에서 키를 붙여 넣은 뒤 **저장(Ctrl+S)** 했는지 확인하세요. "
                    "저장하지 않은 탭 내용은 앱이 읽지 않습니다. 저장 후 브라우저를 새로고침하세요."
                )

    sb: Client | None = None
    if cfg_ok:
        try:
            sb = supabase_client(str(cfg["SUPABASE_URL"]), str(cfg["SUPABASE_ANON_KEY"]))
        except Exception as e:  # noqa: BLE001
            st.error(f"Supabase 연결 실패: {e}")
            sb = None

    render_sidebar_docs(sb, cfg_ok, cfg)

    st.divider()
    col_a, col_b = st.columns((1, 1))

    with col_a:
        genre_key = st.selectbox(
            "장르 선택",
            options=["poetry", "essay", "novel"],
            format_func=lambda k: GENRE_LABELS[k],
        )
        keywords = st.text_area(
            "키워드 / 주제 / 떠오르는 이미지",
            height=120,
            placeholder="예: 빗소리, 늦은 밤 지하철, 할머니의 부엌…",
        )
        mixing = st.radio(
            "문체 믹싱",
            options=("전체 작품 자동 믹싱", "선택한 작품만"),
            horizontal=True,
        )

    selected_doc_ids: list[str] | None = None
    with col_b:
        st.subheader("문체 생성 옵션")
        style_similarity = st.slider(
            "유사도",
            0.0,
            1.0,
            0.12,
            0.01,
            help=(
                "1에 가까울수록 학습 문서와 임베딩이 매우 유사한 청크만 참고해 문체를 강하게 따릅니다. "
                "0에 가까울수록 덜 유사한 청크도 포함되어 창의성·다양성이 커질 수 있습니다. "
                "검색 결과가 없으면 한 번 더 넓게 검색합니다."
            ),
        )

        if cfg_ok and sb is not None:
            try:
                all_docs = list_user_documents(sb)
            except Exception:
                all_docs = []
            if mixing == "선택한 작품만" and all_docs:
                labels = {f"{i + 1}. {d['filename']}": str(d["id"]) for i, d in enumerate(all_docs)}
                picked = st.multiselect("문체로 쓸 문서", options=list(labels.keys()))
                selected_doc_ids = [labels[x] for x in picked] if picked else None
                if not picked:
                    st.caption(
                        "선택한 문서가 없으면 **전체 문서**와 동일하게 검색합니다."
                    )

    st.divider()

    if not cfg_ok or sb is None:
        st.info("키와 DB 연결이 준비되면 중앙 패널에서 창작을 실행할 수 있습니다.")
        return

    oai = openai_client(str(cfg["OPENAI_API_KEY"]))

    if not keywords.strip():
        st.info("**키워드**를 입력한 뒤 **초안 생성**을 누르세요. (키워드가 비어 있으면 생성할 수 없습니다.)")

    novel_plot_text = ""
    if genre_key == "novel":
        st.subheader("소설 기획안")
        st.caption(
            "**PDF, TXT, Markdown(.md), JSON**만 지원합니다. "
            "파일을 선택한 뒤 **「파일 내용을 기획란에 넣기」**를 누르면 아래 입력란에 반영됩니다. "
            "직접 입력·수정도 가능합니다."
        )
        plot_file = st.file_uploader(
            "기획안 파일 업로드",
            type=["pdf", "txt", "md", "json"],
            accept_multiple_files=False,
            key="novel_plot_file_uploader",
        )
        if st.button("파일 내용을 기획란에 넣기", type="secondary", key="novel_plot_apply_btn"):
            if plot_file is None:
                st.warning("먼저 기획안 파일을 선택하세요.")
            else:
                try:
                    raw = plot_file.getvalue()
                    extracted = extract_text_from_upload(plot_file.name, raw).strip()
                    if not extracted:
                        st.error("파일에서 읽을 수 있는 텍스트가 없습니다.")
                    else:
                        st.session_state["novel_plot_editor"] = extracted
                        st.success(f"반영함: {plot_file.name} ({len(extracted):,}자)")
                        st.rerun()
                except ValueError as e:
                    st.error(str(e))
                except Exception as e:  # noqa: BLE001
                    st.error(f"파일 읽기 실패: {e}")

        st.text_area(
            "기획안 (직접 입력·수정)",
            height=220,
            key="novel_plot_editor",
            placeholder="여기에 플롯·등장인물·구성 등을 적거나, 위에서 파일을 불러오세요.",
        )
        novel_plot_text = (st.session_state.get("novel_plot_editor") or "").strip()

    st.subheader("초안 생성")
    if st.button("초안 생성", type="primary"):
        if not keywords.strip():
            st.error("**키워드 / 주제**란에 내용을 입력한 뒤 다시 눌러 주세요.")
            _request_scroll_to_bottom()
        elif genre_key == "novel" and not novel_plot_text:
            st.error("소설은 **기획안**을 입력하거나, 파일을 불러온 뒤 초안을 생성하세요.")
            _request_scroll_to_bottom()
        else:
            try:
                with st.spinner("유사 문체 검색 중…"):
                    rows = match_chunks(
                        sb,
                        oai,
                        keywords.strip(),
                        match_count=STYLE_REFERENCE_TOP_K,
                        match_threshold=float(style_similarity),
                        filter_doc_ids=selected_doc_ids,
                        filter_genre=None,
                    )
                style_block = build_style_block(rows)
            except Exception as e:  # noqa: BLE001
                st.error(str(e))
                _request_scroll_to_bottom()
            else:
                if not style_block.strip():
                    st.error(
                        "**스타일 레퍼런스가 비어 있습니다.** "
                        "사이드바에서 학습 문서를 업로드·처리했는지 확인하세요. "
                        "업로드 후 임베딩이 들어가야 검색됩니다. "
                        "**유사도**를 낮추거나, **문체 믹싱**에서 특정 문서만 고른 경우 선택을 해제해 보세요."
                    )
                    _request_scroll_to_bottom()
                else:
                    try:
                        with st.spinner("초안 작성 중…"):
                            plot_for_draft: str | None = None
                            if genre_key == "novel":
                                plot_for_draft = novel_plot_text or None
                            draft = llm_draft(
                                oai,
                                genre_key,
                                keywords.strip(),
                                style_block,
                                plot_for_draft,
                            )
                            if not draft.strip():
                                st.error(
                                    "모델이 빈 텍스트만 돌려줬습니다. "
                                    "`.env`에 `OPENAI_CHAT_MODEL=gpt-4o-mini`처럼 사용 가능한 모델을 지정해 보세요."
                                )
                                _request_scroll_to_bottom()
                            else:
                                st.session_state["last_draft_saved_at"] = datetime.now(
                                    timezone.utc
                                ).strftime("%Y%m%d_%H%M%S")
                                st.session_state["last_draft"] = draft
                                st.success(
                                    "초안 생성이 완료되었습니다. "
                                    "바로 아래 **생성 결과**로 이동합니다."
                                )
                                _request_scroll_to_bottom()
                    except Exception as e:  # noqa: BLE001
                        st.error(f"초안 작성 오류: {e}")
                        _request_scroll_to_bottom()

    if "last_draft" in st.session_state:
        st.markdown(
            '<div id="hanji-draft-result" style="scroll-margin-top: 4.5rem;"></div>',
            unsafe_allow_html=True,
        )
        st.subheader("생성 결과")
        draft_text = st.session_state["last_draft"]
        stamp = st.session_state.get("last_draft_saved_at")
        file_stem = f"창작초안_{stamp}" if stamp else "창작초안"
        d_col1, d_col2 = st.columns(2)
        with d_col1:
            st.download_button(
                label="Word (.docx) 다운로드",
                data=build_draft_docx_bytes(draft_text),
                file_name=f"{file_stem}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_draft_docx",
                use_container_width=True,
            )
        with d_col2:
            st.download_button(
                label="TXT 다운로드",
                data=draft_text.encode("utf-8-sig"),
                file_name=f"{file_stem}.txt",
                mime="text/plain; charset=utf-8",
                key="download_draft_txt",
                use_container_width=True,
            )
        # 마크다운 특수문자만으로 본문이 사라지는 경우 방지(본문은 그대로 표시)
        st.markdown(draft_text)
        with st.expander("표시가 이상할 때(원문 그대로)", expanded=False):
            st.code(st.session_state["last_draft"], language=None)

    if st.session_state.pop("_pending_scroll_bottom", None):
        _scroll_to_show_new_content()


if __name__ == "__main__":
    main()
